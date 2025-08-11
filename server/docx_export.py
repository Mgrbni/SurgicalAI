"""DOCX export stub for SurgicalAI demo.

Generates a very lightweight .docx using python-docx if installed; otherwise creates a text file with .docx extension.
"""
from __future__ import annotations
from pathlib import Path
from typing import Dict, Any

try:
    from docx import Document  # type: ignore
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


def export_docx(run_dir: Path, context: Dict[str, Any]) -> Path:
    out = run_dir / 'report.docx'
    if DOCX_AVAILABLE:
        doc = Document()
        doc.add_heading('SurgicalAI Analysis Report', 0)
        doc.add_paragraph(f"Run ID: {context.get('run_id')}")
        doc.add_paragraph(f"Primary Diagnosis: {context.get('primary_diagnosis')}")
        doc.add_heading('Top Differentials', level=1)
        for item in context.get('guidelines', []):
            p = doc.add_paragraph()
            p.add_run(f"{item['diagnosis']} ({item['probability']*100:.1f}%): ").bold = True
            if item.get('summary'):
                p.add_run(item['summary'][:300])
        doc.add_heading('Reconstruction', level=1)
        doc.add_paragraph(context.get('reconstruction_text') or 'n/a')
        doc.add_heading('Risk Factors', level=1)
        for k,v in (context.get('risk_factors') or {}).items():
            doc.add_paragraph(f"{k}: {v}")
        doc.add_paragraph('\nGenerated — NOT FOR CLINICAL USE')
        doc.save(out)
    else:
        out.write_text('DOCX dependency missing. Install python-docx for full export.')
    return out
